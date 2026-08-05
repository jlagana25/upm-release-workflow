"""Offline regression tests for album-list document naming migration."""

from __future__ import annotations

import unittest

from config import ReleaseContext
from album_list_doc import _replace_legacy_release_label


class AlbumListDocumentTests(unittest.TestCase):
    def test_legacy_full_label_migrates_outside_album_table_only(self) -> None:
        from docx import Document

        document = Document()
        document.add_paragraph("July 2026 Release")
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "An Album Named July 2026 Release"

        ctx = ReleaseContext(2026, 7, 1, previous_month=True)
        changed = _replace_legacy_release_label(document, ctx)

        self.assertEqual(changed, 1)
        self.assertEqual(document.paragraphs[0].text, "July 2026 (Full) Release")
        self.assertEqual(
            table.cell(0, 0).text,
            "An Album Named July 2026 Release",
        )


if __name__ == "__main__":
    unittest.main()
