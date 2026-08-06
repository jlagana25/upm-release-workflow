"""
folder_setup.py — Steps 2 & 3: Folder Creation
================================================
Step 2  Copy specials baseline → UPM/{specials_root}, replace MMMM YYYY.
Step 3  Copy HD staging + final baselines → dated folders, replace MMMM YYYY,
        copy U-Drive 2.0 User Guide.pdf into MP3 (UDrive 2.0) and WAV (UDrive 2.0).

Prerequisite checks run before any copy:
  • Both Pegasus32 volumes are mounted
  • All three baseline source folders exist

Safety:
  • Existing destination folders abort the step unless --overwrite is set
  • All destructive operations are skipped in --dry-run mode
  • Every file modified, renamed, or copied is logged
  • Failed operations are logged as warnings; the step continues and reports
    overall success/failure at the end

Standalone test:
    python folder_setup.py --test --year 2026 --month 5 --part 1 [--dry-run] [--overwrite]
"""

from __future__ import annotations

import logging
import os
import shutil
from datetime import datetime
from pathlib import Path

from config import (
    BASELINE_HD_FINAL,
    BASELINE_HD_STAGING,
    BASELINE_SPECIALS,
    PLACEHOLDER,
    TEXT_EXTENSIONS,
    VOLUMES,
    ReleaseContext,
    context_from_cli_args,
    is_retired_partner_name,
)


# ---------------------------------------------------------------------------
# Prerequisite check — call before any copy operation
# ---------------------------------------------------------------------------

def verify_prerequisites(
    ctx: ReleaseContext,
    logger: logging.Logger,
) -> bool:
    """
    Confirm required volumes are mounted and all three baseline folders exist.
    Returns True only if every check passes.
    Logs a ✓/✗ line for each item so failures are immediately obvious.
    """
    ok = True

    # Volumes
    for key, vol in VOLUMES.items():
        if vol.exists():
            logger.info(f"  ✓  Volume {key}: {vol}")
        else:
            logger.error(
                f"  ✗  Volume {key} NOT MOUNTED: {vol}\n"
                f"     Mount both Pegasus32 drives before running this step."
            )
            ok = False

    # Baseline source folders
    baselines = {
        "Specials baseline":    BASELINE_SPECIALS,
        "HD Staging baseline":  BASELINE_HD_STAGING,
        "HD Final baseline":    BASELINE_HD_FINAL,
    }
    for label, path in baselines.items():
        if path.exists():
            logger.info(f"  ✓  {label}: {path}")
        else:
            logger.error(
                f"  ✗  {label} NOT FOUND: {path}\n"
                f"     Verify the baseline folder exists before re-running."
            )
            ok = False

    if not ok:
        logger.error(
            "\n  Prerequisite check FAILED.  Resolve the issues above and retry."
        )
    else:
        logger.info("  Prerequisite check passed.")

    return ok


# ---------------------------------------------------------------------------
# Archive helper — non-destructive replacement for shutil.rmtree
# ---------------------------------------------------------------------------

def _archive_existing(
    dst: Path,
    dry_run: bool,
    logger: logging.Logger,
) -> bool:
    """
    Rename dst to '{dst.name}-archived-YYYYMMDD-HHMMSS', preserving all data.

    If a collision occurs (two archives within the same second), appends
    a numeric suffix.  Returns True on success.

    The archive sits next to the original location, so it stays organized
    in the same parent folder as the new copy will.
    """
    timestamp    = datetime.now().strftime("%Y%m%d-%H%M%S")
    archive_path = dst.parent / f"{dst.name}-archived-{timestamp}"

    # Disambiguate in the unlikely case of a same-second collision
    counter = 1
    while archive_path.exists():
        archive_path = dst.parent / f"{dst.name}-archived-{timestamp}-{counter}"
        counter += 1

    if dry_run:
        logger.info(f"  [DRY RUN] Would archive existing →\n    {archive_path}")
        return True

    try:
        dst.rename(archive_path)
        logger.info(f"  Archived existing →\n    {archive_path}")
        return True
    except Exception as exc:
        logger.error(f"  ✗  Could not archive {dst}: {exc}")
        return False


# ---------------------------------------------------------------------------
# Low-level copy helper
# ---------------------------------------------------------------------------

def _is_domo_csv_skeleton(dst: Path) -> bool:
    """
    True if `dst` exists but contains ONLY .csv files (no other file types).

    Step 1 (Domo exports) writes the NBC / Japan / Tunesat metadata CSVs to
    paths *inside* the Specials tree, creating their parent folders via
    ``mkdir(parents=True)``.  So by the time Step 2 runs, a partial
    ``UPM-YYYY-MM-P1/`` skeleton already exists — but it holds only those CSVs,
    not the baseline content.  Treating that as "destination already exists"
    would wrongly block the baseline copy and leave the release tree (e.g. the
    NBC release folder) missing.

    A real prior baseline copy always contains non-CSV files (templates,
    docs, folder structure markers), so "only CSVs present" cleanly
    identifies the Step-1 skeleton case.  An empty dir also counts as a
    skeleton (safe to populate).
    """
    if not dst.exists():
        return False
    for p in dst.rglob("*"):
        if p.is_file() and p.suffix.lower() != ".csv":
            return False
    return True


def _has_unresolved_placeholder_names(dst: Path) -> bool:
    """True when an existing Specials tree is visibly an incomplete copy."""
    if not dst.exists():
        return False
    return any(PLACEHOLDER in path.name for path in dst.rglob("*"))


def _merge_baseline_additive(
    src: Path,
    dst: Path,
    logger: logging.Logger,
) -> int:
    """
    Copy every file from `src` into `dst` that does NOT already exist there,
    creating directories as needed.  Existing files (e.g. the CSVs Step 1
    dropped) are left untouched — this never overwrites.  Returns the number
    of files copied.  Skips *.bat / *.exe to match _safe_copytree's ignore.
    """
    copied = 0
    for root, directories, files in os.walk(src):
        # Do not descend into retired partner trees inherited from the shared
        # baseline. Mutating `directories` is the documented os.walk pruning
        # mechanism and also protects additive/skeleton-resume copies.
        directories[:] = [
            name for name in directories if not is_retired_partner_name(name)
        ]
        rel = os.path.relpath(root, str(src))
        target_dir = dst if rel == "." else dst / rel
        target_dir.mkdir(parents=True, exist_ok=True)
        for fn in files:
            if fn.endswith((".bat", ".exe")):
                continue
            s = Path(root) / fn
            d = target_dir / fn
            if not d.exists():
                shutil.copy2(s, d)
                copied += 1
    return copied


# ---------------------------------------------------------------------------

def _safe_copytree(
    src: Path,
    dst: Path,
    dry_run: bool,
    overwrite: bool,
    label: str,
    logger: logging.Logger,
) -> bool:
    """
    Copy src → dst, respecting dry_run and overwrite flags.

    Returns True if the copy was performed (or would be in dry-run).
    Returns False and logs an error if dst exists and --overwrite is not set.
    """
    if not src.exists():
        logger.error(f"  ✗  Source not found [{label}]: {src}")
        return False

    # Distinguish three states of an existing destination:
    #   (a) doesn't exist            → plain copytree below
    #   (b) Step-1 CSV skeleton only → merge baseline in around the CSVs
    #   (c) real prior copy          → require --overwrite (archive) as before
    merge_into_skeleton = False
    if dst.exists():
        if overwrite:
            logger.warning(
                f"  Destination exists — archiving (--overwrite) [{label}]:\n"
                f"    {dst}"
            )
            if not _archive_existing(dst, dry_run, logger):
                return False
        elif _is_domo_csv_skeleton(dst):
            merge_into_skeleton = True
            logger.info(
                f"  Destination holds only Step-1 metadata CSVs [{label}] — "
                f"merging the baseline in around them (existing CSVs kept):\n"
                f"    {dst}"
            )
        elif _has_unresolved_placeholder_names(dst):
            merge_into_skeleton = True
            logger.warning(
                f"  Destination contains unresolved {PLACEHOLDER!r} names "
                f"[{label}] — treating it as an interrupted baseline copy and "
                f"resuming additively (existing files kept):\n"
                f"    {dst}"
            )
        else:
            logger.error(
                f"  ✗  Destination already exists [{label}]: {dst}\n"
                f"     Pass --overwrite to archive it (renamed with timestamp,\n"
                f"     not deleted), or --skip-folder-setup to bypass this step."
            )
            return False

    if dry_run:
        action = "merge baseline into" if merge_into_skeleton else "copy"
        logger.info(
            f"  [DRY RUN] Would {action} [{label}]:\n"
            f"    {src}\n"
            f"    → {dst}"
        )
    elif merge_into_skeleton:
        logger.info(
            f"  Merging [{label}]:\n"
            f"    {src}\n"
            f"    → {dst}"
        )
        try:
            n = _merge_baseline_additive(src, dst, logger)
        except Exception as exc:
            logger.error(
                f"  ✗ Merge failed [{label}]: {type(exc).__name__}: {exc}"
            )
            return False
        logger.info(f"  Merge complete [{label}]: {n} file(s) added.")
    else:
        logger.info(
            f"  Copying [{label}]:\n"
            f"    {src}\n"
            f"    → {dst}"
        )
        try:
            def _ignore(_directory: str, names: list[str]) -> set[str]:
                return {
                    name
                    for name in names
                    if name.lower().endswith((".bat", ".exe"))
                    or is_retired_partner_name(name)
                }

            shutil.copytree(src, dst, ignore=_ignore)
        except KeyboardInterrupt:
            logger.error(
                f"  ✗ Copy interrupted [{label}] — archiving the partial "
                f"destination so the next run starts clean."
            )
            if dst.exists() and not _archive_existing(
                dst, dry_run=False, logger=logger
            ):
                logger.error(
                    f"     ⚠ Could not archive the partial copy at {dst}.\n"
                    f"     Rename it manually before retrying."
                )
            raise
        except Exception as exc:
            logger.error(
                f"  ✗ Copy failed [{label}]: {type(exc).__name__}: {exc}"
            )
            # A partial destination would block a clean retry (it now
            # "exists") and could be mistaken for a complete copy.  Move it
            # aside so re-running starts fresh, but never silently delete.
            if dst.exists():
                if _archive_existing(dst, dry_run=False, logger=logger):
                    logger.error(
                        "     Partial copy was archived; re-run to retry."
                    )
                else:
                    logger.error(
                        f"     ⚠ Could not archive the partial copy at {dst}.\n"
                        f"     Remove or rename it manually before re-running."
                    )
            return False
        logger.info(f"  Copy complete [{label}].")

    return True


# ---------------------------------------------------------------------------
# Placeholder replacement — text files
# ---------------------------------------------------------------------------

def _replace_in_text_files(
    root: Path,
    replacement: str,
    logger: logging.Logger,
) -> int:
    """
    Replace PLACEHOLDER in all plain-text files under root (in-place).
    Returns the count of files modified.
    """
    count = 0
    for path in root.rglob("*"):
        if not path.is_file():
            continue
        if path.suffix.lower() not in TEXT_EXTENSIONS:
            continue
        try:
            content = path.read_text(encoding="latin-1")
            if PLACEHOLDER in content:
                path.write_text(
                    content.replace(PLACEHOLDER, replacement),
                    encoding="latin-1",
                )
                logger.debug(f"    Text replaced: {path.relative_to(root)}")
                count += 1
        except Exception as exc:
            logger.warning(f"    Could not update {path.relative_to(root)}: {exc}")
    return count


# ---------------------------------------------------------------------------
# Placeholder replacement — DOCX files
# ---------------------------------------------------------------------------

def _replace_in_docx_files(
    root: Path,
    replacement: str,
    logger: logging.Logger,
) -> int:
    """
    Replace PLACEHOLDER in .docx files under root, covering:
      • Body paragraphs
      • Table cells
      • Headers and footers

    Operates at the run level to preserve character formatting.
    Returns the count of files modified.
    """
    try:
        from docx import Document  # type: ignore[import]
    except ImportError:
        logger.warning(
            "  python-docx not installed — skipping .docx placeholder replacement.\n"
            "  Install with:  pip install python-docx"
        )
        return 0

    def _replace_paragraphs(paragraphs) -> bool:
        changed = False
        for para in paragraphs:
            for run in para.runs:
                if PLACEHOLDER in run.text:
                    run.text = run.text.replace(PLACEHOLDER, replacement)
                    changed = True
        return changed

    count = 0
    for path in root.rglob("*.docx"):
        if path.name.startswith("~"):
            continue  # skip Word temporary lock files
        try:
            doc     = Document(str(path))
            changed = False

            # Body paragraphs
            changed |= _replace_paragraphs(doc.paragraphs)

            # Table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        changed |= _replace_paragraphs(cell.paragraphs)

            # Headers and footers in all sections
            for section in doc.sections:
                for hf in (
                    section.header,
                    section.footer,
                    section.even_page_header,
                    section.even_page_footer,
                    section.first_page_header,
                    section.first_page_footer,
                ):
                    try:
                        changed |= _replace_paragraphs(hf.paragraphs)
                    except Exception:
                        pass  # some sections may not have all header/footer types

            if changed:
                doc.save(str(path))
                logger.debug(f"    DOCX replaced: {path.relative_to(root)}")
                count += 1

        except Exception as exc:
            logger.warning(
                f"    Could not update DOCX {path.relative_to(root)}: {exc}"
            )

    return count


# ---------------------------------------------------------------------------
# Placeholder replacement — filesystem names
# ---------------------------------------------------------------------------

def _rename_items(
    root: Path,
    replacement: str,
    logger: logging.Logger,
) -> int:
    """
    Rename files and directories whose names contain PLACEHOLDER.

    Processes deepest paths first (bottom-up) so children are renamed
    before their parents, avoiding stale path references.
    Returns the count of items renamed.
    """
    all_paths = sorted(
        root.rglob("*"),
        key=lambda p: len(p.parts),
        reverse=True,
    )
    count = 0
    for path in all_paths:
        if PLACEHOLDER not in path.name:
            continue
        new_name = path.name.replace(PLACEHOLDER, replacement)
        new_path = path.parent / new_name
        try:
            path.rename(new_path)
            logger.debug(f"    Renamed: {path.name!r} → {new_name!r}")
            count += 1
        except Exception as exc:
            logger.warning(f"    Could not rename {path.name!r}: {exc}")

    return count


# ---------------------------------------------------------------------------
# Replacement orchestrator
# ---------------------------------------------------------------------------

def _apply_placeholder_replacements(
    root: Path,
    replacement: str,
    dry_run: bool,
    logger: logging.Logger,
) -> None:
    """
    Run all three replacement passes under root:
      1. Plain-text file contents
      2. DOCX paragraph/table/header/footer runs
      3. File and directory names (bottom-up)

    In dry-run mode, logs what would be replaced without touching anything.
    """
    if dry_run:
        # Count candidates without modifying
        text_count = sum(
            1 for p in root.rglob("*")
            if p.is_file()
            and p.suffix.lower() in TEXT_EXTENSIONS
            and PLACEHOLDER in p.read_text(encoding="latin-1", errors="ignore")
        ) if root.exists() else 0
        name_count = sum(
            1 for p in root.rglob("*")
            if PLACEHOLDER in p.name
        ) if root.exists() else 0
        logger.info(
            f"  [DRY RUN] Would replace {PLACEHOLDER!r} → {replacement!r} "
            f"in ~{text_count} text file(s) and rename ~{name_count} path(s)"
        )
        return

    logger.info(
        f"  Replacing {PLACEHOLDER!r} → {replacement!r} under {root.name}/"
    )
    t = _replace_in_text_files(root, replacement, logger)
    d = _replace_in_docx_files(root, replacement, logger)
    r = _rename_items(root, replacement, logger)
    logger.info(
        f"  Replacement complete: "
        f"{t} text file(s), {d} DOCX file(s), {r} path(s) renamed"
    )


# ---------------------------------------------------------------------------
# Step 2 — Create Main Specials Folder
# ---------------------------------------------------------------------------

def create_specials_folder(
    ctx: ReleaseContext,
    dry_run: bool,
    overwrite: bool,
    logger: logging.Logger,
) -> bool:
    """
    Copy the UPM Specials Baseline into:
        /Volumes/Pegasus32 R8 - 1/_Specials/UPM/{specials_root}

    Then replace MMMM YYYY throughout the new tree with ctx.month_display_folder.
    """
    logger.info(f"  Source: {BASELINE_SPECIALS}")
    logger.info(f"  Target: {ctx.specials_dir}")

    ok = _safe_copytree(
        BASELINE_SPECIALS, ctx.specials_dir, dry_run, overwrite,
        "Specials", logger,
    )
    if not ok:
        return False

    _apply_placeholder_replacements(
        ctx.specials_dir, ctx.month_display_folder, dry_run, logger
    )

    if not dry_run:
        logger.info(f"  ✓  Specials folder ready: {ctx.specials_dir}")

    return True


# ---------------------------------------------------------------------------
# Step 3 — Create HD Update Folders
# ---------------------------------------------------------------------------

def create_hd_folders(
    ctx: ReleaseContext,
    dry_run: bool,
    overwrite: bool,
    logger: logging.Logger,
) -> bool:
    """
    3a. Copy HD staging baseline → 2-STAGING/{hd_folder}
    3b. Copy HD final baseline   → 3-FINAL PACKAGING/UPM-US/{hd_folder}
    3c. Replace MMMM YYYY in both trees.
    3d. Copy U-Drive 2.0 User Guide.pdf to both UDrive subfolders.

    Returns True only if all four substeps succeed.
    """
    logger.info(f"  HD folder: {ctx.hd_folder}")
    success = True

    # 3a ── HD Staging
    logger.info(f"  3a  Staging: {ctx.hd_staging_dir}")
    ok_staging = _safe_copytree(
        BASELINE_HD_STAGING, ctx.hd_staging_dir, dry_run, overwrite,
        "HD Staging", logger,
    )
    if ok_staging:
        _apply_placeholder_replacements(
            ctx.hd_staging_dir, ctx.month_display_folder, dry_run, logger
        )
    else:
        success = False

    # 3b ── HD Final Packaging
    logger.info(f"  3b  Final:   {ctx.hd_final_dir}")
    ok_final = _safe_copytree(
        BASELINE_HD_FINAL, ctx.hd_final_dir, dry_run, overwrite,
        "HD Final", logger,
    )
    if ok_final:
        _apply_placeholder_replacements(
            ctx.hd_final_dir, ctx.month_display_folder, dry_run, logger
        )
    else:
        success = False

    # 3d ── U-Drive guide
    # Only copy when the HD final tree was freshly created this run.
    # If ok_final is False (destination pre-existed), the guide was already
    # there from a previous successful run — skip to avoid false-positive.
    if dry_run:
        logger.info(
            f"  [DRY RUN] Would copy U-Drive 2.0 User Guide.pdf to:\n"
            f"    {ctx.hd_final_dir / 'MP3 (UDrive 2.0)'}\n"
            f"    {ctx.hd_final_dir / 'WAV (UDrive 2.0)'}"
        )
    elif ok_final:
        _copy_udrive_guide(ctx, logger)
    else:
        logger.info(
            "  Skipping U-Drive guide copy — HD final tree was not created "
            "this run (destination pre-existed or copy failed)."
        )

    if success and not dry_run:
        logger.info(
            f"  ✓  HD staging ready: {ctx.hd_staging_dir}\n"
            f"  ✓  HD final ready:   {ctx.hd_final_dir}"
        )

    return success


def _copy_udrive_guide(ctx: ReleaseContext, logger: logging.Logger) -> None:
    """
    Copy U-Drive 2.0 User Guide.pdf from the HD staging root to both
    MP3 (UDrive 2.0) and WAV (UDrive 2.0) in the HD final dir.

    Non-fatal: logs a warning if the source or either target folder is missing.
    """
    guide_src = ctx.hd_staging_dir / "U-Drive 2.0 User Guide.pdf"

    if not guide_src.exists():
        logger.warning(
            f"  U-Drive 2.0 User Guide.pdf not found at:\n"
            f"    {guide_src}\n"
            f"  Skipping guide copy — confirm the staging baseline contains "
            f"this file."
        )
        return

    targets = [
        ctx.hd_final_dir / "MP3 (UDrive 2.0)",
        ctx.hd_final_dir / "WAV (UDrive 2.0)",
    ]
    for target_dir in targets:
        if not target_dir.exists():
            logger.warning(
                f"  UDrive subfolder not found — was it renamed during "
                f"placeholder replacement?\n    {target_dir}"
            )
            continue
        dest = target_dir / guide_src.name
        try:
            shutil.copy2(guide_src, dest)
            logger.info(f"  Copied U-Drive guide → {dest}")
        except Exception as exc:
            logger.error(
                f"  Failed to copy U-Drive guide to {dest}: {exc}"
            )


# ---------------------------------------------------------------------------
# Standalone test entry point
# ---------------------------------------------------------------------------

def _run_test(args) -> None:
    import sys
    logging.basicConfig(
        level=logging.DEBUG if args.debug else logging.INFO,
        format="%(asctime)s  %(levelname)-8s  %(message)s",
        datefmt="%H:%M:%S",
    )
    logger = logging.getLogger("folder_setup_test")

    ctx = context_from_cli_args(args)
    logger.info(f"Release context: {ctx}")
    logger.info(
        f"  specials_root:        {ctx.specials_root}\n"
        f"  hd_folder:            {ctx.hd_folder}\n"
        f"  month_display_folder: {ctx.month_display_folder}\n"
        f"  dry_run:              {args.dry_run}\n"
        f"  overwrite:            {args.overwrite}"
    )

    # Preflight
    if not verify_prerequisites(ctx, logger):
        sys.exit(1)

    step = args.step.lower()

    if step in ("2", "specials", "all"):
        logger.info("\n── STEP 2: Specials Folder ──")
        ok = create_specials_folder(ctx, args.dry_run, args.overwrite, logger)
        logger.info(f"Step 2: {'OK' if ok else 'FAILED'}")

    if step in ("3", "hd", "all"):
        logger.info("\n── STEP 3: HD Update Folders ──")
        ok = create_hd_folders(ctx, args.dry_run, args.overwrite, logger)
        logger.info(f"Step 3: {'OK' if ok else 'FAILED'}")


if __name__ == "__main__":
    import argparse

    p = argparse.ArgumentParser(description="Test folder creation steps.")
    p.add_argument("--test",     action="store_true", required=True)
    p.add_argument("--year",     type=int)
    p.add_argument("--month",    type=int)
    p.add_argument("--part",     type=int, choices=[1, 2])
    p.add_argument(
        "--previous-month", action="store_true",
        help="Full-month run for the previous month "
             "(no Part split). Relative to today, or to "
             "--year/--month if given.")
    p.add_argument("--step",     default="all",
                   choices=["2", "3", "all", "specials", "hd"],
                   help="Which step to run (default: all)")
    p.add_argument("--dry-run",  action="store_true",
                   help="Log what would happen without making any changes.")
    p.add_argument("--overwrite", action="store_true",
                   help="If destination exists, archive it (rename with "
                        "timestamp suffix, non-destructive) before copying.")
    p.add_argument("--debug",    action="store_true",
                   help="Show DEBUG-level logs (individual file operations).")

    args = p.parse_args()
    _run_test(args)
