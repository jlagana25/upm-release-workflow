"""
album_list_doc.py — Step 4: Album List DOCX and PDF
====================================================
Flow:
  1. Read   /Users/hdfuser/Documents/UPM Tracklists/Release Lists/Album Lists/
              UPM-US-{tracklist_token}-AlbumList.csv
  2. Open   HD-Staging/{hd_folder}/Universal Production Music - {month_display_folder} Album List.docx
            (created by folder_setup.py)
  3. Replace MMMM YYYY → {month_display_folder} in body, tables, headers, footers
  4. Remove any pre-existing tables (idempotent re-runs)
  5. Insert the CSV rows as a Word table
  6. Save the DOCX
  7. Export to PDF at the same location using LibreOffice in headless mode
  8. Copy the PDF into:
        HD-Final/{hd_folder}/MP3 (UDrive 2.0)/
        HD-Final/{hd_folder}/WAV (UDrive 2.0)/

PDF conversion method: LibreOffice `soffice --headless --convert-to pdf`.
This is the most reliable open-source method on macOS and produces
faithful renderings of Word documents.  Install with:
    brew install --cask libreoffice

Fallback location checked: /Applications/LibreOffice.app/Contents/MacOS/soffice

Standalone test:
    python album_list_doc.py --test --year 2026 --month 5 --part 1 [--dry-run]
"""

from __future__ import annotations

import csv
import logging
import shutil
import subprocess
from pathlib import Path

from config import DOCX_TO_PDF_METHODS, PLACEHOLDER, ReleaseContext, context_from_cli_args

# Standard macOS LibreOffice install path (additional to PATH search)
MACOS_SOFFICE = Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")

# Time limit for PDF conversion subprocess call (seconds)
PDF_CONVERT_TIMEOUT = 180

# Built-in Word table style — always available; safe default
DEFAULT_TABLE_STYLE = "Table Grid"

# Number of empty paragraphs to insert between the last content paragraph
# and the table.  One blank line gives the table room to breathe under
# the "<Month> Release" heading.
TABLE_TOP_SPACING_PARAGRAPHS = 1


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def create_album_list_doc(
    ctx: ReleaseContext,
    dry_run: bool,
    logger: logging.Logger,
) -> bool:
    """
    Orchestrate the full Step 4 flow.  Returns True iff every substep
    succeeded (or was correctly skipped in dry-run).
    """
    logger.info(f"  Source CSV:  {ctx.album_list_csv}")
    logger.info(f"  Target DOCX: {ctx.album_list_docx}")
    logger.info(f"  Target PDF:  {ctx.album_list_pdf}")

    # ── Preflight ────────────────────────────────────────────────────────
    # In a dry run, the upstream steps (1 Domo CSV, 3 HD-staging DOCX template)
    # have only been *described*, not actually produced — so their outputs may
    # legitimately not exist yet.  Treat missing prerequisites as a dry-run
    # no-op (warn + succeed) rather than a hard failure, so a full-pipeline
    # --dry-run can preview past this step.  In a real run they remain errors.
    if not ctx.album_list_csv.exists():
        msg = (
            f"  Album list CSV not found: {ctx.album_list_csv}\n"
            f"     Run Step 1 (Domo exports) first."
        )
        if dry_run:
            logger.warning("  ⚠ " + msg.strip())
            logger.info(
                "  [DRY RUN] Skipping album-list generation preview "
                "(prerequisite CSV not present yet)."
            )
            return True
        logger.error("  ✗ " + msg)
        return False

    if not ctx.album_list_docx.exists():
        msg = (
            f"  DOCX template not found: {ctx.album_list_docx}\n"
            f"     Run Step 3 (folder_setup) first to create HD staging."
        )
        if dry_run:
            logger.warning("  ⚠ " + msg.strip())
            logger.info(
                "  [DRY RUN] Skipping album-list generation preview "
                "(DOCX template not present yet)."
            )
            return True
        logger.error("  ✗ " + msg)
        return False

    # ── 1. Read CSV ──────────────────────────────────────────────────────
    rows = _read_csv(ctx.album_list_csv, logger)
    if rows is None:
        return False
    if not rows:
        logger.error("  ✗  CSV is empty — refusing to produce a blank album list.")
        return False

    # ── 2–5. Update DOCX (placeholder + table) and save ──────────────────
    if not _build_docx(ctx, rows, dry_run, logger):
        return False

    # ── 6. Export PDF ────────────────────────────────────────────────────
    if not _export_pdf(ctx, dry_run, logger):
        return False

    # ── 7. Copy PDF into both UDrive folders ─────────────────────────────
    if not _distribute_pdf(ctx, dry_run, logger):
        return False

    if dry_run:
        logger.info("  [DRY RUN] All Step 4 operations would succeed.")
    else:
        logger.info(f"  ✓  Album list ready: {ctx.album_list_pdf}")
    return True


# ---------------------------------------------------------------------------
# CSV reading
# ---------------------------------------------------------------------------

def _read_csv(
    csv_path: Path,
    logger: logging.Logger,
) -> list[dict] | None:
    """
    Read the album list CSV into a list of OrderedDict rows.
    Returns None on read failure.
    """
    try:
        with open(csv_path, encoding="utf-8-sig", newline="") as f:
            reader = csv.DictReader(f)
            rows = list(reader)
        logger.info(
            f"  Read {len(rows)} row(s) from CSV.  "
            f"Columns: {list(rows[0].keys()) if rows else 'none'}"
        )
        return rows
    except Exception as exc:
        logger.error(f"  ✗  Could not read CSV: {exc}")
        return None


# ---------------------------------------------------------------------------
# DOCX manipulation
# ---------------------------------------------------------------------------

def _build_docx(
    ctx: ReleaseContext,
    rows: list[dict],
    dry_run: bool,
    logger: logging.Logger,
) -> bool:
    """Replace placeholders, remove old tables, insert new table, save."""
    if dry_run:
        logger.info(
            f"  [DRY RUN] Would replace {PLACEHOLDER!r} → "
            f"{ctx.month_display_text!r} and insert a "
            f"{len(rows)}-row table into {ctx.album_list_docx.name}"
        )
        return True

    try:
        from docx import Document  # type: ignore[import]
    except ImportError:
        logger.error(
            "  ✗  python-docx not installed.\n"
            "     Run:  pip install python-docx"
        )
        return False

    try:
        doc = Document(str(ctx.album_list_docx))

        n_replaced = _replace_placeholder_in_doc(
            doc, PLACEHOLDER, ctx.month_display_text
        )
        logger.info(f"  Replaced {n_replaced} placeholder run(s).")
        if n_replaced == 0:
            n_migrated = _replace_legacy_release_label(doc, ctx)
            if n_migrated:
                logger.info(
                    f"  Migrated {n_migrated} legacy release-label run(s) "
                    f"to {ctx.month_display_text!r}."
                )

        n_removed = _remove_existing_tables(doc)
        if n_removed:
            logger.info(
                f"  Removed {n_removed} pre-existing table(s) "
                f"for idempotent re-runs."
            )

        # Remove trailing empty paragraphs so the table sits near the top
        # content instead of being pushed to the bottom of the page.
        n_blanks = _remove_trailing_empty_paragraphs(doc)
        if n_blanks:
            logger.info(
                f"  Removed {n_blanks} trailing empty paragraph(s) "
                f"to lift the table up the page."
            )

        _insert_csv_table(doc, rows, logger)

        doc.save(str(ctx.album_list_docx))
        logger.info(f"  DOCX saved: {ctx.album_list_docx}")
        return True

    except Exception as exc:
        logger.error(f"  ✗  DOCX build failed: {exc}")
        return False


def _replace_placeholder_in_doc(doc, placeholder: str, replacement: str) -> int:
    """
    Replace `placeholder` with `replacement` across the whole document,
    including placeholders that Word has split across multiple runs
    (a common occurrence with copy-paste, autocorrect, or formatting marks).

    Covers body paragraphs, table cells, and all header/footer variants.
    Returns the count of replacements performed.
    """
    count = 0

    def _replace_paragraphs(paragraphs) -> int:
        n = 0
        for para in paragraphs:
            n += _replace_in_paragraph(para, placeholder, replacement)
        return n

    count += _replace_paragraphs(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                count += _replace_paragraphs(cell.paragraphs)

    for section in doc.sections:
        for hf in (
            section.header,
            section.footer,
            section.even_page_header,
            section.even_page_footer,
            section.first_page_header,
            section.first_page_footer,
        ):
            try:
                count += _replace_paragraphs(hf.paragraphs)
            except Exception:
                pass

    return count


def _replace_legacy_release_label(doc, ctx: ReleaseContext) -> int:
    """Upgrade an already-generated document to explicit Part/Full wording.

    Once Step 4 has consumed ``MMMM YYYY``, a later naming migration cannot use
    the placeholder path again.  Restrict this fallback to body paragraphs and
    headers/footers (never album-table cells) so a legitimate album title that
    contains the month is not altered.
    """
    if ctx.month_display_text == ctx.month_display:
        return 0
    legacy = f"{ctx.month_display} Release"
    replacement = f"{ctx.month_display_text} Release"
    count = 0
    for paragraph in doc.paragraphs:
        count += _replace_in_paragraph(paragraph, legacy, replacement)
    for section in doc.sections:
        for hf in (
            section.header,
            section.footer,
            section.even_page_header,
            section.even_page_footer,
            section.first_page_header,
            section.first_page_footer,
        ):
            try:
                for paragraph in hf.paragraphs:
                    count += _replace_in_paragraph(
                        paragraph, legacy, replacement
                    )
            except Exception:
                pass
    return count


def _replace_in_paragraph(paragraph, placeholder: str, replacement: str) -> int:
    """
    Replace every occurrence of placeholder in a single paragraph, handling
    the case where Word has split the placeholder across multiple runs.

    Strategy: locate the placeholder in the concatenated text, then update
    the affected runs:
      - If the placeholder lies entirely within one run, replace it there.
      - If it spans runs, put the replacement in the first affected run,
        clear any fully-consumed middle runs, and trim the leading-consumed
        portion off the last affected run.

    Preserves the formatting of the first affected run for the replacement
    text — which is the standard behaviour for run-level docx editing.
    """
    count = 0
    while placeholder in paragraph.text:
        runs = list(paragraph.runs)
        if not runs:
            break

        # Map each run to (start, end) offsets in the concatenated text
        positions = []
        offset = 0
        for run in runs:
            positions.append((run, offset, offset + len(run.text)))
            offset += len(run.text)

        full_text = "".join(r.text for r in runs)
        ph_start  = full_text.find(placeholder)
        if ph_start < 0:
            break
        ph_end = ph_start + len(placeholder)

        # Which runs overlap [ph_start, ph_end)?
        affected = [
            i for i, (_, s, e) in enumerate(positions)
            if s < ph_end and e > ph_start
        ]
        if not affected:
            break

        first_i, last_i = affected[0], affected[-1]
        first_run, first_s, _ = positions[first_i]
        last_run,  last_s,  _ = positions[last_i]

        # Text in the first run before the placeholder begins
        before = first_run.text[:ph_start - first_s]
        # Text in the last run after the placeholder ends
        after  = last_run.text[ph_end - last_s:]

        if first_i == last_i:
            # Placeholder fits inside a single run
            first_run.text = before + replacement + after
        else:
            # Spans multiple runs
            first_run.text = before + replacement
            # Clear the fully-consumed middle runs
            for i in affected[1:-1]:
                positions[i][0].text = ""
            # Last run keeps only what was AFTER the placeholder
            last_run.text = after

        count += 1

    return count


def _remove_trailing_empty_paragraphs(doc) -> int:
    """
    Delete empty (whitespace-only) paragraphs from the end of the body.
    These cause the inserted table to be pushed far down the page.
    Stops at the first paragraph that contains real content.
    Returns the count of paragraphs removed.
    """
    count = 0
    paragraphs = list(doc.paragraphs)
    for para in reversed(paragraphs):
        if para.text.strip() == "":
            element = para._element
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)
                count += 1
        else:
            break
    return count


def _remove_existing_tables(doc) -> int:
    """
    Delete every table currently in the document body.
    Returns the count removed.

    Necessary so re-runs don't append duplicate tables.  The baseline
    DOCX is text-only, so this is non-destructive in normal operation.
    """
    count = 0
    # iterate over a list copy because we're mutating during traversal
    for table in list(doc.tables):
        element = table._element
        element.getparent().remove(element)
        count += 1
    return count


def _insert_csv_table(
    doc,
    rows: list[dict],
    logger: logging.Logger,
) -> None:
    """Append a new Word table at the end of the document with all CSV rows."""
    headers = list(rows[0].keys())

    # Insert one or more blank paragraphs to separate the table from the
    # preceding heading text — without these the table butts up against
    # "<Month> Release".
    for _ in range(TABLE_TOP_SPACING_PARAGRAPHS):
        doc.add_paragraph("")

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = DEFAULT_TABLE_STYLE
    except KeyError:
        # Built-in styles only exist if the document has them defined.
        # Falling back silently is fine — table still renders.
        pass

    # Header row
    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = header
        # Bold the header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    # Data rows
    for row_idx, data_row in enumerate(rows, start=1):
        for col_idx, header in enumerate(headers):
            table.rows[row_idx].cells[col_idx].text = str(
                data_row.get(header, "") or ""
            )

    logger.info(
        f"  Inserted table: {len(rows)} row(s) × {len(headers)} column(s)."
    )


# ---------------------------------------------------------------------------
# PDF export
# ---------------------------------------------------------------------------

def _find_soffice(logger: logging.Logger) -> str | None:
    """
    Locate LibreOffice's `soffice` binary.  Checks (in order):
      1. PATH using config.DOCX_TO_PDF_METHODS
      2. /Applications/LibreOffice.app/Contents/MacOS/soffice  (macOS app)
    """
    for name in DOCX_TO_PDF_METHODS:
        which = shutil.which(name)
        if which:
            logger.debug(f"  PDF converter found on PATH: {which}")
            return which

    if MACOS_SOFFICE.exists():
        logger.debug(f"  PDF converter found at: {MACOS_SOFFICE}")
        return str(MACOS_SOFFICE)

    logger.error(
        "  ✗  Could not find LibreOffice.\n"
        "     Install with:  brew install --cask libreoffice\n"
        "     Or download from https://www.libreoffice.org/"
    )
    return None


def _export_pdf(
    ctx: ReleaseContext,
    dry_run: bool,
    logger: logging.Logger,
) -> bool:
    """Convert the DOCX to PDF using LibreOffice headless mode."""
    if dry_run:
        logger.info(f"  [DRY RUN] Would export PDF: {ctx.album_list_pdf}")
        return True

    soffice = _find_soffice(logger)
    if not soffice:
        return False

    # Delete pre-existing PDF so we know the new one really got produced
    if ctx.album_list_pdf.exists():
        try:
            ctx.album_list_pdf.unlink()
            logger.debug(f"  Removed stale PDF: {ctx.album_list_pdf.name}")
        except Exception as exc:
            logger.warning(f"  Could not remove stale PDF: {exc}")

    cmd = [
        soffice,
        "--headless",
        "--convert-to", "pdf",
        "--outdir", str(ctx.album_list_docx.parent),
        str(ctx.album_list_docx),
    ]
    logger.info("  Running LibreOffice headless PDF conversion…")
    try:
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=PDF_CONVERT_TIMEOUT,
        )
    except subprocess.TimeoutExpired:
        logger.error(
            f"  ✗  PDF conversion timed out after {PDF_CONVERT_TIMEOUT}s."
        )
        return False
    except Exception as exc:
        logger.error(f"  ✗  PDF conversion failed to launch: {exc}")
        return False

    if result.returncode != 0:
        logger.error(
            f"  ✗  soffice exited with code {result.returncode}\n"
            f"     stderr: {result.stderr.strip()[-500:]}"
        )
        return False

    if not ctx.album_list_pdf.exists():
        logger.error(
            f"  ✗  PDF was not produced at expected path:\n"
            f"     {ctx.album_list_pdf}\n"
            f"     soffice stdout: {result.stdout.strip()[-500:]}"
        )
        return False

    logger.info(f"  PDF created: {ctx.album_list_pdf}")
    return True


# ---------------------------------------------------------------------------
# PDF distribution
# ---------------------------------------------------------------------------

def _distribute_pdf(
    ctx: ReleaseContext,
    dry_run: bool,
    logger: logging.Logger,
) -> bool:
    """Copy the PDF to MP3 (UDrive 2.0) and WAV (UDrive 2.0) folders."""
    targets = [
        ctx.hd_final_dir / "MP3 (UDrive 2.0)",
        ctx.hd_final_dir / "WAV (UDrive 2.0)",
    ]

    success = True
    for target_dir in targets:
        if dry_run:
            logger.info(
                f"  [DRY RUN] Would copy PDF → "
                f"{target_dir / ctx.album_list_pdf.name}"
            )
            continue

        if not target_dir.exists():
            logger.error(
                f"  ✗  Target folder missing: {target_dir}\n"
                f"     Run Step 3 (folder_setup) to create HD final folders."
            )
            success = False
            continue

        dest = target_dir / ctx.album_list_pdf.name
        try:
            shutil.copy2(ctx.album_list_pdf, dest)
            logger.info(f"  Copied PDF → {dest}")
        except Exception as exc:
            logger.error(f"  ✗  Failed to copy PDF to {dest}: {exc}")
            success = False

    return success


# ---------------------------------------------------------------------------
# Standalone test entry point
# ---------------------------------------------------------------------------

def _run_test(args) -> None:
    import sys
    logging.basicConfig(
        level=logging.DEBUG if args.debug else logging.INFO,
        format="%(asctime)s  %(levelname)-8s  %(message)s",
        datefmt="%H:%M:%S",
    )
    logger = logging.getLogger("album_list_doc_test")

    ctx = context_from_cli_args(args)
    logger.info(f"Release context: {ctx}")
    logger.info(f"  dry_run: {args.dry_run}")

    ok = create_album_list_doc(ctx, args.dry_run, logger)
    sys.exit(0 if ok else 1)


if __name__ == "__main__":
    import argparse

    p = argparse.ArgumentParser(
        description="Generate the Album List DOCX and PDF."
    )
    p.add_argument("--test",    action="store_true", required=True)
    p.add_argument("--year",    type=int)
    p.add_argument("--month",   type=int)
    p.add_argument("--part",    type=int, choices=[1, 2])
    p.add_argument(
        "--previous-month", action="store_true",
        help="Full-month run for the previous month "
             "(no Part split). Relative to today, or to "
             "--year/--month if given.")
    p.add_argument("--dry-run", action="store_true",
                   help="Log what would happen without making changes.")
    p.add_argument("--debug",   action="store_true",
                   help="Verbose logging.")

    args = p.parse_args()
    _run_test(args)
